from docx import Document
from num2words import num2words
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import comtypes.client
import pythoncom
import os


def format_date(date_str):
    try:
        return datetime.strptime(date_str, "%Y-%m-%d").strftime("%d.%m.%Y")
    except ValueError:
        return date_str


def set_cell_borders(cell, border_color="000000", border_size="4"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    borders = OxmlElement('w:tcBorders')
    for border_name in ('top', 'left', 'bottom', 'right'):
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), border_size)
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), border_color)
        borders.append(border)
    tcPr.append(borders)


def generate_document_from_template(data, template_path):
    doc = Document(template_path)

    data['act_date'] = format_date(data['act_date'])
    data['invoice_date'] = format_date(data['invoice_date'])
    for service in data['services']:
        service['date'] = format_date(service['date'])

    for para in doc.paragraphs:
        for run in para.runs:
            run.text = run.text.replace("{ACT_NUMBER}", str(data['act_number']))
            run.text = run.text.replace("{ACT_DATE}", str(data['act_date']))
            run.text = run.text.replace("{CLIENT_NAME}", str(data['client_name']))
            run.text = run.text.replace("{CLIENT_ADDRESS}", str(data['client_address']))
            run.text = run.text.replace("{CLIENT_PHONE}", str(data['client_phone']))
            run.text = run.text.replace("{CLIENT_EMAIL}", str(data['client_email']))
            run.text = run.text.replace("{INVOICE_NUMBER}", str(data['invoice_number']))
            run.text = run.text.replace("{INVOICE_DATE}", str(data['invoice_date']))

            run.text = run.text.replace("{EXECUTOR_NAME}", str(data['executor_name']))
            run.text = run.text.replace("{EXECUTOR_ADDRESS}", str(data['executor_address']))
            run.text = run.text.replace("{EXECUTOR_PHONE}", str(data['executor_phone']))
            run.text = run.text.replace("{EXECUTOR_EMAIL}", str(data['executor_email']))
            run.text = run.text.replace("{EXECUTOR_SIGNATURE}", str(data['executor_signature']))

            run.text = run.text.replace("{CLIENT_SIGNATURE}", str(data['client_signature']))

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.text = run.text.replace("{ACT_NUMBER}", str(data['act_number']))
                        run.text = run.text.replace("{ACT_DATE}", str(data['act_date']))
                        run.text = run.text.replace("{CLIENT_NAME}", str(data['client_name']))
                        run.text = run.text.replace("{CLIENT_ADDRESS}", str(data['client_address']))
                        run.text = run.text.replace("{CLIENT_PHONE}", str(data['client_phone']))
                        run.text = run.text.replace("{CLIENT_EMAIL}", str(data['client_email']))
                        run.text = run.text.replace("{INVOICE_NUMBER}", str(data['invoice_number']))
                        run.text = run.text.replace("{INVOICE_DATE}", str(data['invoice_date']))

                        run.text = run.text.replace("{EXECUTOR_NAME}", str(data['executor_name']))
                        run.text = run.text.replace("{EXECUTOR_ADDRESS}", str(data['executor_address']))
                        run.text = run.text.replace("{EXECUTOR_PHONE}", str(data['executor_phone']))
                        run.text = run.text.replace("{EXECUTOR_EMAIL}", str(data['executor_email']))
                        run.text = run.text.replace("{EXECUTOR_SIGNATURE}", str(data['executor_signature']))

                        run.text = run.text.replace("{CLIENT_SIGNATURE}", str(data['client_signature']))

    for table in doc.tables:
        for i, row in enumerate(table.rows):
            if any("{NUM_SERVICE}" in cell.text for cell in row.cells):
                template_index = i
                template_row = row
                break
        else:
            continue

        for index, service in enumerate(data['services']):
            if index == 0:
                cells = template_row.cells
            else:
                new_row = table.add_row()
                cells = new_row.cells
                for j, template_cell in enumerate(template_row.cells):
                    new_paragraph = cells[j].paragraphs[0]
                    for run in template_cell.paragraphs[0].runs:
                        new_run = new_paragraph.add_run(run.text)
                        new_run.bold = run.bold
                        new_run.italic = run.italic
                        new_run.font.name = run.font.name
                        new_run.font.size = run.font.size

            cells[0].text = str(index + 1)
            cells[1].text = service['date']
            cells[2].text = service['car_number']
            cells[3].text = service['service_name']
            cells[4].text = str(service['quantity'])
            cells[5].text = str(service['price'])
            cells[6].text = str(service['total'])

            for cell in cells:
                set_cell_borders(cell)
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.font.size = Pt(10)

        if "{NUM_SERVICE}" in template_row.cells[0].text:
            table._tbl.remove(template_row._element)

    total_sum = sum(service['total'] for service in data['services'])
    total_sum_text = num2words(total_sum, lang='ru').capitalize() + ' рублей 00 копеек'

    for para in doc.paragraphs:
        for run in para.runs:
            run.text = run.text.replace("{TOTAL_SUM}", f"{total_sum:.2f}")
            run.text = run.text.replace("{TOTAL_SUM_TEXT}", total_sum_text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.text = run.text.replace("{TOTAL_SUM}", f"{total_sum:.2f}")
                        run.text = run.text.replace("{TOTAL_SUM_TEXT}", total_sum_text)

    # Очистка имени файла от недопустимых символов
    def sanitize_filename(filename):
        # Удаляем все символы, кроме букв, цифр, пробелов, дефисов и точек
        return "".join(c for c in filename if c.isalnum() or c in " ._-()№")

    # Генерация имени файла без недопустимых символов
    filename = sanitize_filename(f"Акт №{data['act_number']} для {data['client_name']} от {data['act_date']}.docx")
    word_path = os.path.join("static", "docs", filename)

    # Сохраняем Word-документ
    doc.save(word_path)

    # Генерация пути к PDF
    pdf_path = word_path.replace('.docx', '.pdf')


    pythoncom.CoInitialize()
    try:
        word = comtypes.client.CreateObject("Word.Application")
        doc = word.Documents.Open(os.path.abspath(word_path))
        doc.SaveAs(os.path.abspath(pdf_path), FileFormat=17)
        doc.Close()
        word.Quit()
    finally:
        pythoncom.CoUninitialize()

    return word_path, pdf_path
